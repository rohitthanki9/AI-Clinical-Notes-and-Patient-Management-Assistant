"""
Export Manager for AI Clinical Notes Assistant
Exports clinical notes to PDF and DOCX formats
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from datetime import datetime
from pathlib import Path

class ExportManager:
    def __init__(self, clinic_name="Medical Clinic"):
        """
        Initialize export manager
        Args:
            clinic_name: Name of clinic for headers
        """
        self.clinic_name = clinic_name

    def export_to_docx(self, content, output_path, note_type="Clinical Note",
                       patient_info=None, doctor_info=None, logo_path=None):
        """
        Export note to DOCX format
        Args:
            content: Note content
            output_path: Path to save DOCX file
            note_type: Type of note
            patient_info: Patient information dict
            doctor_info: Doctor information dict
            logo_path: Path to clinic logo
        Returns:
            True if successful, False otherwise
        """
        try:
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)

            # Add logo if available
            if logo_path and Path(logo_path).exists():
                doc.add_picture(logo_path, width=Inches(1.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add clinic name
            clinic_heading = doc.add_heading(self.clinic_name, level=1)
            clinic_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add note type
            note_heading = doc.add_heading(note_type, level=2)
            note_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Spacing

            # Add date
            date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.add_paragraph()  # Spacing

            # Add patient information
            if patient_info:
                doc.add_heading("Patient Information", level=3)
                patient_table = doc.add_table(rows=0, cols=2)
                patient_table.style = 'Light Grid Accent 1'

                info_items = [
                    ("Name", patient_info.get('name', 'N/A')),
                    ("Age", f"{patient_info.get('age', 'N/A')} years" if patient_info.get('age') else 'N/A'),
                    ("Gender", patient_info.get('gender', 'N/A')),
                    ("Contact", patient_info.get('contact', 'N/A'))
                ]

                for label, value in info_items:
                    row = patient_table.add_row()
                    row.cells[0].text = label
                    row.cells[1].text = str(value)

                doc.add_paragraph()  # Spacing

            # Add doctor information
            if doctor_info:
                doc.add_heading("Provider Information", level=3)
                doc.add_paragraph(f"Doctor: Dr. {doctor_info.get('name', 'Unknown')}")
                if doctor_info.get('email'):
                    doc.add_paragraph(f"Email: {doctor_info.get('email')}")

                doc.add_paragraph()  # Spacing

            # Add horizontal line
            doc.add_paragraph('_' * 80)
            doc.add_paragraph()

            # Add note content
            # Split content into sections and format
            for line in content.split('\n'):
                if line.strip():
                    # Check if it's a heading (ends with :)
                    if line.strip().endswith(':') and len(line.strip()) < 50:
                        doc.add_heading(line.strip(), level=3)
                    else:
                        doc.add_paragraph(line)
                else:
                    doc.add_paragraph()  # Preserve spacing

            # Save document
            doc.save(str(output_path))
            return True

        except Exception as e:
            print(f"Error exporting to DOCX: {e}")
            return False

    def export_to_pdf(self, content, output_path, note_type="Clinical Note",
                      patient_info=None, doctor_info=None, logo_path=None):
        """
        Export note to PDF format
        Args:
            content: Note content
            output_path: Path to save PDF file
            note_type: Type of note
            patient_info: Patient information dict
            doctor_info: Doctor information dict
            logo_path: Path to clinic logo
        Returns:
            True if successful, False otherwise
        """
        try:
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )

            # Container for PDF elements
            elements = []

            # Define styles
            styles = getSampleStyleSheet()

            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=RGBColor(0, 51, 102),
                spaceAfter=12,
                alignment=TA_CENTER
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=RGBColor(0, 51, 102),
                spaceAfter=6,
                spaceBefore=12
            )

            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6
            )

            # Add logo if available
            if logo_path and Path(logo_path).exists():
                try:
                    logo = Image(logo_path, width=1.5*inch, height=1.5*inch)
                    elements.append(logo)
                    elements.append(Spacer(1, 0.2*inch))
                except:
                    pass

            # Add clinic name
            elements.append(Paragraph(self.clinic_name, title_style))
            elements.append(Spacer(1, 0.1*inch))

            # Add note type
            elements.append(Paragraph(note_type, heading_style))
            elements.append(Spacer(1, 0.2*inch))

            # Add date
            date_text = f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            elements.append(Paragraph(date_text, normal_style))
            elements.append(Spacer(1, 0.2*inch))

            # Add patient information
            if patient_info:
                elements.append(Paragraph("Patient Information", heading_style))
                elements.append(Paragraph(f"<b>Name:</b> {patient_info.get('name', 'N/A')}", normal_style))
                if patient_info.get('age'):
                    elements.append(Paragraph(f"<b>Age:</b> {patient_info.get('age')} years", normal_style))
                if patient_info.get('gender'):
                    elements.append(Paragraph(f"<b>Gender:</b> {patient_info.get('gender')}", normal_style))
                if patient_info.get('contact'):
                    elements.append(Paragraph(f"<b>Contact:</b> {patient_info.get('contact')}", normal_style))
                elements.append(Spacer(1, 0.2*inch))

            # Add doctor information
            if doctor_info:
                elements.append(Paragraph("Provider Information", heading_style))
                elements.append(Paragraph(f"<b>Doctor:</b> Dr. {doctor_info.get('name', 'Unknown')}", normal_style))
                if doctor_info.get('email'):
                    elements.append(Paragraph(f"<b>Email:</b> {doctor_info.get('email')}", normal_style))
                elements.append(Spacer(1, 0.2*inch))

            # Add horizontal line
            elements.append(Paragraph("_" * 80, normal_style))
            elements.append(Spacer(1, 0.2*inch))

            # Add note content
            for line in content.split('\n'):
                if line.strip():
                    # Check if it's a heading
                    if line.strip().endswith(':') and len(line.strip()) < 50:
                        elements.append(Paragraph(f"<b>{line.strip()}</b>", heading_style))
                    else:
                        # Escape special characters for ReportLab
                        safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        elements.append(Paragraph(safe_line, normal_style))
                else:
                    elements.append(Spacer(1, 0.1*inch))

            # Build PDF
            doc.build(elements)
            return True

        except Exception as e:
            print(f"Error exporting to PDF: {e}")
            return False

    def export_encrypted(self, content, output_path, encryption_manager,
                        format='pdf', **kwargs):
        """
        Export note and encrypt the file
        Args:
            content: Note content
            output_path: Path to save file
            encryption_manager: EncryptionManager instance
            format: 'pdf' or 'docx'
            **kwargs: Additional arguments for export functions
        Returns:
            Path to encrypted file if successful, None otherwise
        """
        try:
            # Export to temporary file
            temp_path = Path(output_path).with_suffix(f'.{format}')

            if format == 'pdf':
                success = self.export_to_pdf(content, temp_path, **kwargs)
            elif format == 'docx':
                success = self.export_to_docx(content, temp_path, **kwargs)
            else:
                return None

            if not success:
                return None

            # Encrypt the file
            encrypted_path = encryption_manager.encrypt_file(temp_path)

            # Remove unencrypted temporary file
            if temp_path.exists():
                temp_path.unlink()

            return encrypted_path

        except Exception as e:
            print(f"Error in encrypted export: {e}")
            return None
